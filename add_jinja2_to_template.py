#!/usr/bin/env python3
"""
在復興自主學習申請表.docx中新增Jinja2變數標記
"""

from docx import Document
import shutil
import os

# 路徑設定
input_path = "./復興自主學習申請表.docx"
output_path = "./附件一 復興自主學習申請表-新版.docx"

# 先複製原始檔案
shutil.copy(input_path, output_path)
print(f"已複製: {input_path} -> {output_path}")

# 開啟檔案
doc = Document(output_path)
table = doc.tables[0]

print(f"表格共有 {len(table.rows)} 行")
print("\n開始新增Jinja2變數標記...\n")

try:
    # 第2行: 計畫名稱 (列1開始，跨越多列)
    # 找到計畫名稱的輸入格
    title_row = table.rows[2]
    # 在第1列之後的儲存格中新增 {{title}}
    for i in range(1, len(title_row.cells)):
        cell = title_row.cells[i]
        if cell.text.strip() == "" or "是否繳交" not in cell.text:
            # 清空並新增變數
            for para in cell.paragraphs:
                para.clear()
            cell.paragraphs[0].add_run("{{title}}")
            print(f"✓ 第2行列{i}: 新增 {{{{title}}}}")
            break

    # 第4行: 組長資訊
    member1_row = table.rows[4]
    # 班級 (列1-2)
    if len(member1_row.cells) > 2:
        cell = member1_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        cell.paragraphs[0].add_run("{{member1.class_name}}")
        print("✓ 第4行: 組長班級")

    # 座號 (列3)
    if len(member1_row.cells) > 3:
        cell = member1_row.cells[3]
        for para in cell.paragraphs:
            para.clear()
        cell.paragraphs[0].add_run("{{member1.seat_number}}")
        print("✓ 第4行: 組長座號")

    # 學號 (列4-5)
    if len(member1_row.cells) > 4:
        cell = member1_row.cells[4]
        for para in cell.paragraphs:
            para.clear()
        cell.paragraphs[0].add_run("{{member1.student_id}}")
        print("✓ 第4行: 組長學號")

    # 第5行: 組員2資訊
    member2_row = table.rows[5]
    if len(member2_row.cells) > 1:
        cell = member2_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        cell.paragraphs[0].add_run("{{member2.class_name}}")
    if len(member2_row.cells) > 3:
        cell = member2_row.cells[3]
        for para in cell.paragraphs:
            para.clear()
        cell.paragraphs[0].add_run("{{member2.seat_number}}")
    if len(member2_row.cells) > 4:
        cell = member2_row.cells[4]
        for para in cell.paragraphs:
            para.clear()
        cell.paragraphs[0].add_run("{{member2.student_id}}")
    print("✓ 第5行: 組員2資訊")

    # 第6行: 組員3資訊
    member3_row = table.rows[6]
    if len(member3_row.cells) > 1:
        cell = member3_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        cell.paragraphs[0].add_run("{{member3.class_name}}")
    if len(member3_row.cells) > 3:
        cell = member3_row.cells[3]
        for para in cell.paragraphs:
            para.clear()
        cell.paragraphs[0].add_run("{{member3.seat_number}}")
    if len(member3_row.cells) > 4:
        cell = member3_row.cells[4]
        for para in cell.paragraphs:
            para.clear()
        cell.paragraphs[0].add_run("{{member3.student_id}}")
    print("✓ 第6行: 組員3資訊")

    # 第7行: 學習動機 (列1，跨越整行)
    motivation_row = table.rows[7]
    if len(motivation_row.cells) > 1:
        cell = motivation_row.cells[1]
        # 清空所有段落
        for para in cell.paragraphs:
            para.clear()
        # 如果沒有段落，新增一個
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        cell.paragraphs[0].add_run("{{motivation}}")
        print("✓ 第7行: 學習動機")

    # 第8行: 學習類別
    category_row = table.rows[8]
    if len(category_row.cells) > 1:
        cell = category_row.cells[1]
        # 保留原有文字，在checkbox處新增條件
        original_text = cell.text
        for para in cell.paragraphs:
            para.clear()
        # 新增帶條件的選項
        para = cell.paragraphs[0]
        text = (
            "{% if learning_categories.get('閱讀計畫') %}☑{% else %}□{% endif %}閱讀計畫 "
            "{% if learning_categories.get('志工服務') %}☑{% else %}□{% endif %}志工服務 "
            "{% if learning_categories.get('專題研究') %}☑{% else %}□{% endif %}專題研究 "
            "{% if learning_categories.get('藝文創作') %}☑{% else %}□{% endif %}藝文創作 "
            "{% if learning_categories.get('技藝學習') %}☑{% else %}□{% endif %}技藝學習 "
            "{% if learning_categories.get('競賽準備') %}☑{% else %}□{% endif %}競賽準備\n"
            "{% if learning_categories.get('實作體驗') %}☑{% else %}□{% endif %}實作體驗 "
            "{% if learning_categories.get('課程延伸') %}☑{% else %}□{% endif %}課程延伸 "
            "□其他：{{learning_category_other}}"
        )
        para.add_run(text)
        print("✓ 第8行: 學習類別")

    # 第9行: 學習環境需求
    env_row = table.rows[9]
    if len(env_row.cells) > 1:
        cell = env_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        para = cell.paragraphs[0]
        text = (
            "A. 圖書館場地:\n"
            "{% if env_needs.get('自習室') %}☑{% else %}□{% endif %}自習室(低噪音，可自備筆電者) "
            "{% if env_needs.get('數位閱讀室') %}☑{% else %}□{% endif %}數位閱讀室(需電腦操作者)\n"
            "{% if env_needs.get('雲端教室') %}☑{% else %}□{% endif %}雲端教室(需小組討論者) "
            "{% if env_needs.get('美力教室') %}☑{% else %}□{% endif %}美力教室(有實作或大桌面需求者)\n"
            "B. 其他場地：{{env_other}}"
        )
        para.add_run(text)
        print("✓ 第9行: 學習環境需求")

    # 第10行: 學習裝置需求
    equipment_row = table.rows[10]
    for i in range(1, len(equipment_row.cells)):
        cell = equipment_row.cells[i]
        if cell.text.strip() == "" or i == 1:
            for para in cell.paragraphs:
                para.clear()
            cell.paragraphs[0].add_run("{{equipment_needs}}")
            print(f"✓ 第10行列{i}: 學習裝置需求")
            break

    # 第11行: 學習方法/參考資料
    reference_row = table.rows[11]
    if len(reference_row.cells) > 1:
        cell = reference_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        text = "{% for ref in references %}{{ref.book_title}} / {{ref.author}} / {{ref.publisher}}{% if ref.link %} / {{ref.link}}{% endif %}{% if not loop.last %}\n{% endif %}{% endfor %}"
        cell.paragraphs[0].add_run(text)
        print("✓ 第11行: 參考資料")

    # 第12行: 預期成效
    outcome_row = table.rows[12]
    if len(outcome_row.cells) > 1:
        cell = outcome_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        cell.paragraphs[0].add_run("{{expected_outcome}}")
        print("✓ 第12行: 預期成效")

    # 第14-22行: 學習內容規劃 (9個項次)
    print("\n新增學習內容規劃...")
    for idx in range(9):
        row_num = 14 + idx
        if row_num < len(table.rows):
            row = table.rows[row_num]
            # 日期 (列2-3)
            if len(row.cells) > 2:
                cell = row.cells[2]
                for para in cell.paragraphs:
                    para.clear()
                cell.paragraphs[0].add_run(f"{{% if plan_items|length > {idx} %}}{{{{plan_items[{idx}].date}}}}{{% endif %}}")

            # 時數 (列4)
            if len(row.cells) > 4:
                cell = row.cells[4]
                for para in cell.paragraphs:
                    para.clear()
                cell.paragraphs[0].add_run(f"{{% if plan_items|length > {idx} %}}{{{{plan_items[{idx}].hours}}}}{{% endif %}}")

            # 學習內容 (列5-6)
            if len(row.cells) > 5:
                cell = row.cells[5]
                for para in cell.paragraphs:
                    para.clear()
                cell.paragraphs[0].add_run(f"{{% if plan_items|length > {idx} %}}{{{{plan_items[{idx}].content}}}}{{% endif %}}")

            # 檢核指標 (列7-9)
            if len(row.cells) > 7:
                cell = row.cells[7]
                for para in cell.paragraphs:
                    para.clear()
                cell.paragraphs[0].add_run(f"{{% if plan_items|length > {idx} %}}{{{{plan_items[{idx}].metric}}}}{{% endif %}}")

            print(f"  ✓ 項次 {idx + 1}")

    # 第23行: 階段中目標
    midterm_row = table.rows[23]
    if len(midterm_row.cells) > 1:
        cell = midterm_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        cell.paragraphs[0].add_run("{{midterm_goal}}")
        print("\n✓ 第23行: 階段中目標")

    # 第24行: 階段末目標
    final_row = table.rows[24]
    if len(final_row.cells) > 1:
        cell = final_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        cell.paragraphs[0].add_run("{{final_goal}}")
        print("✓ 第24行: 階段末目標")

    # 第25行: 成果發表形式
    presentation_row = table.rows[25]
    if len(presentation_row.cells) > 1:
        cell = presentation_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        para = cell.paragraphs[0]
        text = (
            "{% if '靜態展' in (presentation_format or '') %}☑{% else %}□{% endif %}靜態展(PPT、書面報告、心得、自我省思…等)\n"
            "{% if '動態展' in (presentation_format or '') %}☑{% else %}□{% endif %}動態展(直播、影片撥放、實際展示、演出…等)\n"
            "□其他 {{presentation_other}}"
        )
        para.add_run(text)
        print("✓ 第25行: 成果發表形式")

    # 儲存檔案
    doc.save(output_path)
    print(f"\n✅ 成功! 已儲存到: {output_path}")
    print("\n下一步: 將此檔案複製到Docker容器中")
    print("執行: docker cp \"附件一 復興自主學習申請表-新版.docx\" self-learning-backend:/app/templates/application_template.docx")

except Exception as e:
    print(f"\n❌ 錯誤: {e}")
    import traceback
    traceback.print_exc()
