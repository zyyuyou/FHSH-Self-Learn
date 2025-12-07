#!/usr/bin/env python3
"""
正確地新增參考資料行 - 使用table.add_row()然後手動設定每個cell
"""
from docx import Document

template_path = "./附件一 復興自主學習申請表-新版.docx"

print(f"讀取模板: {template_path}")
doc = Document(template_path)

table = doc.tables[0]

# 首先，更新第10行為標題行
print("\n=== 更新第10行為標題行 ===")
row_10 = table.rows[10]

# 設定標題
row_10.cells[0].text = "學習方法\n(參考資料)"
row_10.cells[1].text = "書名"
row_10.cells[2].text = ""  # 書名延伸
row_10.cells[3].text = ""  # 書名延伸
row_10.cells[4].text = "作者"
row_10.cells[5].text = ""  # 作者延伸
row_10.cells[6].text = "出版社"
row_10.cells[7].text = ""  # 出版社延伸
row_10.cells[8].text = "連結（電子書籍或電子期刊）"
row_10.cells[9].text = ""  # 連結延伸

print("  已設定標題行")

# 在第11行之前插入3個新行（參考資料資料行）
# 我們需要在Row 10之後的位置插入
print("\n=== 插入3個參考資料資料行 ===")

# 獲取第11行的索引
row_11_idx = 11

# 插入3個參考資料行（從後往前插入，保持順序）
for ref_idx in range(2, -1, -1):  # 2, 1, 0
    print(f"  插入參考資料 {ref_idx + 1} 的行...")

    # 在row_11_idx位置插入新行
    # 注意：python-docx的table._tbl.insert需要訪問底層XML
    # 更簡單的方法：在表格末尾新增行，然後移動它

    # 新增新行到表格末尾
    new_row = table.add_row()

    # 設定Cell 0（與標題合並）
    new_row.cells[0].text = ""

    # Cell 1-3: 書名
    for i in [1, 2, 3]:
        new_row.cells[i].text = ""
        new_row.cells[i].paragraphs[0].text = f"{{% if references|length > {ref_idx} %}}{{{{references[{ref_idx}].book_title}}}}{{% endif %}}"

    # Cell 4-5: 作者
    for i in [4, 5]:
        new_row.cells[i].text = ""
        new_row.cells[i].paragraphs[0].text = f"{{% if references|length > {ref_idx} %}}{{{{references[{ref_idx}].author}}}}{{% endif %}}"

    # Cell 6-7: 出版社
    for i in [6, 7]:
        new_row.cells[i].text = ""
        new_row.cells[i].paragraphs[0].text = f"{{% if references|length > {ref_idx} %}}{{{{references[{ref_idx}].publisher}}}}{{% endif %}}"

    # Cell 8-9: 連結
    for i in [8, 9]:
        new_row.cells[i].text = ""
        new_row.cells[i].paragraphs[0].text = f"{{% if references|length > {ref_idx} %}}{{{{references[{ref_idx}].link}}}}{{% endif %}}"

    # 將這個新行移動到正確的位置（第11行）
    # 獲取表格的XML元素
    tbl = table._tbl
    # 獲取新行的XML元素（最後一行）
    new_row_xml = tbl[-1]
    # 獲取目標位置的XML元素（第11行 = 索引10）
    target_row_xml = tbl[row_11_idx]
    # 移除最後一行
    tbl.remove(new_row_xml)
    # 在目標位置之前插入
    target_row_xml.addprevious(new_row_xml)

print("\n✅ 完成！")

doc.save(template_path)
print(f"\n已儲存: {template_path}")
