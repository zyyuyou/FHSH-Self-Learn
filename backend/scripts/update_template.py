#!/usr/bin/env python3
"""
更新 Word 模板，新增 Jinja2 變數標記
"""
from docx import Document
from docx.shared import Pt, RGBColor
from pathlib import Path

def update_template():
    template_path = Path("/app/templates/application_template.docx")

    # 載入原始模板
    doc = Document(template_path)

    print(f"原始模板有 {len(doc.paragraphs)} 個段落")
    print(f"原始模板有 {len(doc.tables)} 個表格")

    # 顯示前幾個段落以瞭解結構
    print("\n前30個段落:")
    for i, para in enumerate(doc.paragraphs[:30]):
        if para.text.strip():
            print(f"{i}: {para.text}")

    # 如果有表格，顯示第一個表格的結構
    if doc.tables:
        print("\n\n第一個表格的前10行:")
        table = doc.tables[0]
        for row_idx, row in enumerate(table.rows[:15]):
            cells_text = [cell.text.strip() for cell in row.cells]
            print(f"Row {row_idx}: {cells_text}")

    # 儲存以便檢查
    output_path = Path("/tmp/template_analysis.txt")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(f"總段落數: {len(doc.paragraphs)}\n")
        f.write(f"總表格數: {len(doc.tables)}\n\n")

        f.write("所有段落:\n")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                f.write(f"{i}: {para.text}\n")

        if doc.tables:
            f.write("\n\n所有表格:\n")
            for table_idx, table in enumerate(doc.tables):
                f.write(f"\n表格 {table_idx}:\n")
                for row_idx, row in enumerate(table.rows):
                    cells_text = [cell.text.strip() for cell in row.cells]
                    f.write(f"  Row {row_idx}: {cells_text}\n")

    print(f"\n已將分析結果儲存到: {output_path}")

if __name__ == "__main__":
    update_template()
