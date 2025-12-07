#!/usr/bin/env python3
"""
更新 Word 模板中的參考資料欄位
將參考資料改為支援多個專案（類似學習內容規劃的做法）
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

template_path = "./附件一 復興自主學習申請表-新版.docx"
output_path = "./附件一 復興自主學習申請表-新版.docx"

print(f"讀取模板: {template_path}")
doc = Document(template_path)

# 獲取第一個表格
table = doc.tables[0]

# 找到第10行（學習方法(參考資料)）
row_10 = table.rows[10]

print("\n更新第10行（學習方法(參考資料)）的內容...")
print(f"Row 10 有 {len(row_10.cells)} 個 cells")

# 檢查cells是否合併
print("\n檢查cell合併狀態...")
for i in range(min(10, len(row_10.cells))):
    cell = row_10.cells[i]
    print(f"Cell {i}: {len(cell.paragraphs)} 段落")

# 定義參考資料的模板 - 將所有參考資料合併到一個欄位顯示
# 使用類似學習內容規劃的方式，用換行分隔不同的參考資料
cell_1 = row_10.cells[1]
# 清空所有段落
for p in cell_1.paragraphs:
    p.clear()

# 新增模板變數 - 顯示所有參考資料，每個一行
if len(cell_1.paragraphs) > 0:
    para = cell_1.paragraphs[0]
else:
    para = cell_1.add_paragraph()

# 使用Jinja2迴圈來渲染所有參考資料
template_text = """{% for ref in references %}{{ ref.book_title }} / {{ ref.author }} / {{ ref.publisher }}{% if ref.link %} / {{ ref.link }}{% endif %}
{% endfor %}"""
para.text = template_text

print(f"  已更新 Cell 1 為參考資料模板")

print(f"\n儲存更新後的模板: {output_path}")
doc.save(output_path)
print("✅ 完成！")
