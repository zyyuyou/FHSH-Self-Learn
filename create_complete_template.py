#!/usr/bin/env python3
"""
完整的模板生成指令碼 - 包含所有欄位和正確的參考資料處理
"""

from docx import Document
import shutil

# 路徑設定
input_path = "./復興自主學習申請表.docx"
output_path = "./附件一 復興自主學習申請表-新版.docx"

# 先複製原始檔案
shutil.copy(input_path, output_path)
print(f"已複製: {input_path} -> {output_path}")

# 開啟檔案
doc = Document(output_path)
table = doc.tables[0]

print(f"\n開始建立完整模板...")
print("=" * 80)

try:
    # ===== 第1行: 計畫名稱 =====
    title_row = table.rows[1]
    if len(title_row.cells) > 1:
        cell = title_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        cell.paragraphs[0].add_run("{{title}}")
        print("✓ 第1行: 計畫名稱")

    # ===== 第3行: 學生1（組長）資料 =====
    member1_row = table.rows[3]
    if len(member1_row.cells) > 1:
        cell = member1_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        cell.paragraphs[0].add_run("{{member1.class_name}}")

    if len(member1_row.cells) > 3:
        cell = member1_row.cells[3]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        cell.paragraphs[0].add_run("{{member1.seat_number}}")

    if len(member1_row.cells) > 4:
        cell = member1_row.cells[4]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        cell.paragraphs[0].add_run("{{member1.student_id}}")
    print("✓ 第3行: 組長資訊")

    # ===== 第4行: 學生2資料 =====
    member2_row = table.rows[4]
    if len(member2_row.cells) > 1:
        cell = member2_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        cell.paragraphs[0].add_run("{{member2.class_name}}")
    if len(member2_row.cells) > 3:
        cell = member2_row.cells[3]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        cell.paragraphs[0].add_run("{{member2.seat_number}}")
    if len(member2_row.cells) > 4:
        cell = member2_row.cells[4]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        cell.paragraphs[0].add_run("{{member2.student_id}}")
    print("✓ 第4行: 組員2資訊")

    # ===== 第5行: 學生3資料 =====
    member3_row = table.rows[5]
    if len(member3_row.cells) > 1:
        cell = member3_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        cell.paragraphs[0].add_run("{{member3.class_name}}")
    if len(member3_row.cells) > 3:
        cell = member3_row.cells[3]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        cell.paragraphs[0].add_run("{{member3.seat_number}}")
    if len(member3_row.cells) > 4:
        cell = member3_row.cells[4]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        cell.paragraphs[0].add_run("{{member3.student_id}}")
    print("✓ 第5行: 組員3資訊")

    # ===== 第6行: 學習動機 =====
    motivation_row = table.rows[6]
    if len(motivation_row.cells) > 1:
        cell = motivation_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        cell.paragraphs[0].add_run("{{motivation}}")
        print("✓ 第6行: 學習動機")

    # ===== 第7行: 學習類別 =====
    category_row = table.rows[7]
    if len(category_row.cells) > 1:
        cell = category_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        para = cell.paragraphs[0]
        text = (
            "{% if learning_categories.get('閱讀計畫') %}☑{% else %}□{% endif %}閱讀計畫 "
            "{% if learning_categories.get('志工服務') %}☑{% else %}□{% endif %}志工服務 "
            "{% if learning_categories.get('專題研究') %}☑{% else %}□{% endif %}專題研究 "
            "{% if learning_categories.get('藝文創作') %}☑{% else %}□{% endif %}藝文創作 "
            "{% if learning_categories.get('技藝學習') %}☑{% else %}□{% endif %}技藝學習 "
            "{% if learning_categories.get('競賽準備') %}☑{% else %}□{% endif %}競賽準備\n"
            "{% if learning_categories.get('實作體驗') %}☑{% else %}□{% endif %}實作體驗 "
            "{% if learning_categories.get('課程延伸') %}☑{% else %}□{% endif %}課程延伸 "
            "□其他：{{learning_category_other}}"
        )
        para.add_run(text)
        print("✓ 第7行: 學習類別")

    # ===== 第8行: 學習環境需求 =====
    env_row = table.rows[8]
    if len(env_row.cells) > 1:
        cell = env_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        para = cell.paragraphs[0]
        text = (
            "A. 圖書館場地:\n"
            "{% if env_needs.get('自習室') %}☑{% else %}□{% endif %}自習室(低噪音，可自備筆電者) "
            "{% if env_needs.get('數位閱讀室') %}☑{% else %}□{% endif %}數位閱讀室(需電腦操作者)\n"
            "{% if env_needs.get('雲端教室') %}☑{% else %}□{% endif %}雲端教室(需小組討論者) "
            "{% if env_needs.get('美力教室') %}☑{% else %}□{% endif %}美力教室(有實作或大桌面需求者)\n"
            "B. 其他場地：{{env_other}}"
        )
        para.add_run(text)
        print("✓ 第8行: 學習環境需求")

    # ===== 第9行: 學習裝置需求 =====
    equipment_row = table.rows[9]
    if len(equipment_row.cells) > 1:
        cell = equipment_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        cell.paragraphs[0].add_run("{{equipment_needs}}")
        print("✓ 第9行: 學習裝置需求")

    # ===== 第10行: 學習方法/參考資料（內嵌表格） =====
    ref_row = table.rows[10]
    ref_cell = ref_row.cells[1]

    if len(ref_cell.tables) > 0:
        inner_table = ref_cell.tables[0]

        # 內嵌表格第1行是資料填入區
        if len(inner_table.rows) > 1:
            data_row = inner_table.rows[1]

            # 書名（列0）- 包含連結
            cell_book = data_row.cells[0]
            for para in cell_book.paragraphs:
                para.clear()
            if len(cell_book.paragraphs) == 0:
                cell_book.add_paragraph()
            cell_book.paragraphs[0].add_run("{% for ref in references %}{{ref.book_title}}{% if ref.link %}\n連結: {{ref.link}}{% endif %}{% if not loop.last %}\n\n{% endif %}{% endfor %}")

            # 作者（列1）
            cell_author = data_row.cells[1]
            for para in cell_author.paragraphs:
                para.clear()
            if len(cell_author.paragraphs) == 0:
                cell_author.add_paragraph()
            cell_author.paragraphs[0].add_run("{% for ref in references %}{{ref.author}}{% if not loop.last %}\n\n{% endif %}{% endfor %}")

            # 出版社（列2）
            cell_publisher = data_row.cells[2]
            for para in cell_publisher.paragraphs:
                para.clear()
            if len(cell_publisher.paragraphs) == 0:
                cell_publisher.add_paragraph()
            cell_publisher.paragraphs[0].add_run("{% for ref in references %}{{ref.publisher}}{% if not loop.last %}\n\n{% endif %}{% endfor %}")

            print("✓ 第10行: 參考資料（內嵌表格）")
    else:
        print("⚠️  第10行: 未找到內嵌表格")

    # ===== 第11行: 預期成效 =====
    outcome_row = table.rows[11]
    if len(outcome_row.cells) > 1:
        cell = outcome_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        cell.paragraphs[0].add_run("{{expected_outcome}}")
        print("✓ 第11行: 預期成效")

    # ===== 第13-21行: 學習內容規劃 (9個項次) =====
    print("\n新增學習內容規劃...")
    for idx in range(9):
        row_num = 13 + idx
        if row_num < len(table.rows):
            row = table.rows[row_num]

            # 日期 (列2)
            if len(row.cells) > 2:
                cell = row.cells[2]
                for para in cell.paragraphs:
                    para.clear()
                if len(cell.paragraphs) == 0:
                    cell.add_paragraph()
                cell.paragraphs[0].add_run(f"{{% if plan_items|length > {idx} %}}{{{{plan_items[{idx}].date}}}}{{% endif %}}")

            # 時數 (列4)
            if len(row.cells) > 4:
                cell = row.cells[4]
                for para in cell.paragraphs:
                    para.clear()
                if len(cell.paragraphs) == 0:
                    cell.add_paragraph()
                cell.paragraphs[0].add_run(f"{{% if plan_items|length > {idx} %}}{{{{plan_items[{idx}].hours}}}}{{% endif %}}")

            # 學習內容 (列5)
            if len(row.cells) > 5:
                cell = row.cells[5]
                for para in cell.paragraphs:
                    para.clear()
                if len(cell.paragraphs) == 0:
                    cell.add_paragraph()
                cell.paragraphs[0].add_run(f"{{% if plan_items|length > {idx} %}}{{{{plan_items[{idx}].content}}}}{{% endif %}}")

            # 檢核指標 (列7)
            if len(row.cells) > 7:
                cell = row.cells[7]
                for para in cell.paragraphs:
                    para.clear()
                if len(cell.paragraphs) == 0:
                    cell.add_paragraph()
                cell.paragraphs[0].add_run(f"{{% if plan_items|length > {idx} %}}{{{{plan_items[{idx}].metric}}}}{{% endif %}}")

            print(f"  ✓ 項次 {idx + 1}")

    # ===== 第22行: 階段中目標 =====
    midterm_row = table.rows[22]
    if len(midterm_row.cells) > 1:
        cell = midterm_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        cell.paragraphs[0].add_run("{{midterm_goal}}")
        print("\n✓ 第22行: 階段中目標")

    # ===== 第23行: 階段末目標 =====
    final_row = table.rows[23]
    if len(final_row.cells) > 1:
        cell = final_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        cell.paragraphs[0].add_run("{{final_goal}}")
        print("✓ 第23行: 階段末目標")

    # ===== 第24行: 成果發表形式 =====
    presentation_row = table.rows[24]
    if len(presentation_row.cells) > 1:
        cell = presentation_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        para = cell.paragraphs[0]
        text = (
            "{% if '靜態展' in (presentation_format or '') %}☑{% else %}□{% endif %}靜態展(PPT、書面報告、心得、自我省思…等)\n"
            "{% if '動態展' in (presentation_format or '') %}☑{% else %}□{% endif %}動態展(直播、影片撥放、實際展示、演出…等)\n"
            "□其他 {{presentation_other}}"
        )
        para.add_run(text)
        print("✓ 第24行: 成果發表形式")

    # 儲存檔案
    doc.save(output_path)
    print(f"\n{'=' * 80}")
    print(f"✅ 完整模板建立成功！")
    print(f"已儲存到: {output_path}")
    print(f"\n重點修正：")
    print(f"  1. ✅ 計畫名稱在第1行（正確位置）")
    print(f"  2. ✅ 班級標題在第2行保持不變")
    print(f"  3. ✅ 參考資料使用內嵌表格的三個欄位（書名/作者/出版社）")
    print(f"  4. ✅ 所有其他欄位正確配置")
    print(f"\n下一步: 複製到Docker容器")
    print(f'docker cp "附件一 復興自主學習申請表-新版.docx" self-learning-backend:/app/templates/application_template.docx')

except Exception as e:
    print(f"\n❌ 錯誤: {e}")
    import traceback
    traceback.print_exc()
