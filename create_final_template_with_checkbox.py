#!/usr/bin/env python3
"""
最終模板 - 使用checkbox格式的「是否繳交過」且新增姓名欄位
所有填入欄位統一使用「新細明體」字型
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn
import shutil

# 路徑設定
input_path = "./復興自主學習申請表-已取消合併.docx"
output_path = "./附件一 復興自主學習申請表-新版.docx"

def set_run_font(run, font_name="新細明體", font_size=12):
    """
    設定 run 的字型
    :param run: docx.text.run.Run 物件
    :param font_name: 字型名稱
    :param font_size: 字型大小（pt）
    """
    run.font.name = font_name
    run.font.size = Pt(font_size)

    # 設定東亞字型（對中文很重要）
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

# 先複製原始檔案
shutil.copy(input_path, output_path)
print(f"已複製: {input_path} -> {output_path}")

# 開啟檔案
doc = Document(output_path)

# 刪除「□全組人員都已填寫Google表單」段落
print(f"\n刪除 Google 表單提示文字...")
if len(doc.paragraphs) > 0 and 'Google' in doc.paragraphs[0].text:
    # 清空第一個段落的內容
    p = doc.paragraphs[0]
    p.clear()
    print("✓ 已刪除 Google 表單提示")

table = doc.tables[0]

print(f"\n開始建立最終模板（checkbox+姓名欄位）...")
print("=" * 80)

try:
    # ===== 第1行: 計畫名稱 =====
    title_row = table.rows[1]
    if len(title_row.cells) > 1:
        cell = title_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        run = cell.paragraphs[0].add_run("{{title}}")
        set_run_font(run)
        print("✓ 第1行: 計畫名稱")

    # ===== 第3-5行: 學生資料（含姓名） =====
    for idx, student_name in [(3, 'member1'), (4, 'member2'), (5, 'member3')]:
        row = table.rows[idx]

        # 班級
        if len(row.cells) > 1:
            cell = row.cells[1]
            for para in cell.paragraphs:
                para.clear()
            if len(cell.paragraphs) == 0:
                cell.add_paragraph()
            run = cell.paragraphs[0].add_run(f"{{{{{student_name}.class_name}}}}")
            set_run_font(run)

        # 座號
        if len(row.cells) > 3:
            cell = row.cells[3]
            for para in cell.paragraphs:
                para.clear()
            if len(cell.paragraphs) == 0:
                cell.add_paragraph()
            run = cell.paragraphs[0].add_run(f"{{{{{student_name}.seat_number}}}}")
            set_run_font(run)

        # 學號
        if len(row.cells) > 4:
            cell = row.cells[4]
            for para in cell.paragraphs:
                para.clear()
            if len(cell.paragraphs) == 0:
                cell.add_paragraph()
            run = cell.paragraphs[0].add_run(f"{{{{{student_name}.student_id}}}}")
            set_run_font(run)

        # 姓名 - 在列6（cells[6], [7], [8] 可能是合併儲存格，只需處理 cell 6）
        if len(row.cells) > 6:
            cell = row.cells[6]
            # 清空所有段落（包括格式）
            for para in cell.paragraphs:
                para.clear()
            # 如果沒有段落，新增一個
            if len(cell.paragraphs) == 0:
                cell.add_paragraph()

            # 重置段落對齊方式為 None（確保與學生2、3一致）
            cell.paragraphs[0].alignment = None

            # 新增姓名變數
            run = cell.paragraphs[0].add_run(f"{{{{{student_name}.name}}}}")
            set_run_font(run)
            # 注意：cells[7] 和 cells[8] 與 cell[6] 共享底層元素，不需要單獨處理

        # 是否繳交過 - 改為checkbox格式（「是」或「否」兩個選項）
        if len(row.cells) > 9:
            cell = row.cells[9]
            for para in cell.paragraphs:
                para.clear()
            if len(cell.paragraphs) == 0:
                cell.add_paragraph()
            text = (
                f"{{% if {student_name}.has_submitted == '是' %}}☑{{% else %}}□{{% endif %}}是\n"
                f"{{% if {student_name}.has_submitted == '否' %}}☑{{% else %}}□{{% endif %}}否"
            )
            run = cell.paragraphs[0].add_run(text)
            set_run_font(run)

        print(f"✓ 第{idx}行: {student_name}資訊（含姓名和checkbox）")

    # ===== 第6行: 學習動機 =====
    motivation_row = table.rows[6]
    if len(motivation_row.cells) > 1:
        cell = motivation_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        run = cell.paragraphs[0].add_run("{{motivation}}")
        set_run_font(run)
        print("✓ 第6行: 學習動機")

    # ===== 第7行: 學習類別 =====
    category_row = table.rows[7]
    if len(category_row.cells) > 1:
        cell = category_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        para = cell.paragraphs[0]
        text = (
            "{% if learning_categories.get('閱讀計畫') %}☑{% else %}□{% endif %}閱讀計畫 "
            "{% if learning_categories.get('志工服務') %}☑{% else %}□{% endif %}志工服務 "
            "{% if learning_categories.get('專題研究') %}☑{% else %}□{% endif %}專題研究 "
            "{% if learning_categories.get('藝文創作') %}☑{% else %}□{% endif %}藝文創作 "
            "{% if learning_categories.get('技藝學習') %}☑{% else %}□{% endif %}技藝學習 "
            "{% if learning_categories.get('競賽準備') %}☑{% else %}□{% endif %}競賽準備\n"
            "{% if learning_categories.get('實作體驗') %}☑{% else %}□{% endif %}實作體驗 "
            "{% if learning_categories.get('課程延伸') %}☑{% else %}□{% endif %}課程延伸 "
            "{% if learning_category_other %}☑{% else %}□{% endif %}其他：{{learning_category_other}}"
        )
        run = para.add_run(text)
        set_run_font(run)
        print("✓ 第7行: 學習類別")

    # ===== 第8行: 學習環境需求 =====
    env_row = table.rows[8]
    if len(env_row.cells) > 1:
        cell = env_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        para = cell.paragraphs[0]
        text = (
            "A. 圖書館場地:\n"
            "{% if env_needs.get('自習室') %}☑{% else %}□{% endif %}自習室(低噪音，可自備筆電者) "
            "{% if env_needs.get('數位閱讀室') %}☑{% else %}□{% endif %}數位閱讀室(需電腦操作者)\n"
            "{% if env_needs.get('雲端教室') %}☑{% else %}□{% endif %}雲端教室(需小組討論者) "
            "{% if env_needs.get('美力教室') %}☑{% else %}□{% endif %}美力教室(有實作或大桌面需求者)\n"
            "B. 其他場地：{{env_other}}"
        )
        run = para.add_run(text)
        set_run_font(run)
        print("✓ 第8行: 學習環境需求")

    # ===== 第9行: 學習裝置需求 =====
    equipment_row = table.rows[9]
    if len(equipment_row.cells) > 1:
        cell = equipment_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        run = cell.paragraphs[0].add_run("{{equipment_needs}}")
        set_run_font(run)
        print("✓ 第9行: 學習裝置需求")

    # ===== 第10行: 學習方法/參考資料（內嵌表格） =====
    ref_row = table.rows[10]
    ref_cell = ref_row.cells[1]

    if len(ref_cell.tables) > 0:
        inner_table = ref_cell.tables[0]
        if len(inner_table.rows) > 1:
            data_row = inner_table.rows[1]

            # 書名
            cell_book = data_row.cells[0]
            for para in cell_book.paragraphs:
                para.clear()
            if len(cell_book.paragraphs) == 0:
                cell_book.add_paragraph()
            cell_book.paragraphs[0].add_run("{% for ref in references %}{{ref.book_title}}{% if ref.link %}\n連結: {{ref.link}}{% endif %}{% if not loop.last %}\n\n{% endif %}{% endfor %}")

            # 作者
            cell_author = data_row.cells[1]
            for para in cell_author.paragraphs:
                para.clear()
            if len(cell_author.paragraphs) == 0:
                cell_author.add_paragraph()
            cell_author.paragraphs[0].add_run("{% for ref in references %}{{ref.author}}{% if not loop.last %}\n\n{% endif %}{% endfor %}")

            # 出版社
            cell_publisher = data_row.cells[2]
            for para in cell_publisher.paragraphs:
                para.clear()
            if len(cell_publisher.paragraphs) == 0:
                cell_publisher.add_paragraph()
            cell_publisher.paragraphs[0].add_run("{% for ref in references %}{{ref.publisher}}{% if not loop.last %}\n\n{% endif %}{% endfor %}")

            print("✓ 第10行: 參考資料（內嵌表格）")

    # ===== 第11行: 預期成效 =====
    outcome_row = table.rows[11]
    if len(outcome_row.cells) > 1:
        cell = outcome_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        run = cell.paragraphs[0].add_run("{{expected_outcome}}")
        set_run_font(run)
        print("✓ 第11行: 預期成效")

    # ===== 第13-21行: 學習內容規劃 =====
    print("\n新增學習內容規劃...")
    for idx in range(9):
        row_num = 13 + idx
        if row_num < len(table.rows):
            row = table.rows[row_num]

            # 日期
            if len(row.cells) > 2:
                cell = row.cells[2]
                for para in cell.paragraphs:
                    para.clear()
                if len(cell.paragraphs) == 0:
                    cell.add_paragraph()
                run = cell.paragraphs[0].add_run(f"{{% if plan_items|length > {idx} %}}{{{{plan_items[{idx}].date}}}}{{% endif %}}")
                set_run_font(run)

            # 時數
            if len(row.cells) > 4:
                cell = row.cells[4]
                for para in cell.paragraphs:
                    para.clear()
                if len(cell.paragraphs) == 0:
                    cell.add_paragraph()
                run = cell.paragraphs[0].add_run(f"{{% if plan_items|length > {idx} %}}{{{{plan_items[{idx}].hours}}}}{{% endif %}}")
                set_run_font(run)

            # 學習內容
            if len(row.cells) > 5:
                cell = row.cells[5]
                for para in cell.paragraphs:
                    para.clear()
                if len(cell.paragraphs) == 0:
                    cell.add_paragraph()
                run = cell.paragraphs[0].add_run(f"{{% if plan_items|length > {idx} %}}{{{{plan_items[{idx}].content}}}}{{% endif %}}")
                set_run_font(run)

            # 檢核指標
            if len(row.cells) > 7:
                cell = row.cells[7]
                for para in cell.paragraphs:
                    para.clear()
                if len(cell.paragraphs) == 0:
                    cell.add_paragraph()
                run = cell.paragraphs[0].add_run(f"{{% if plan_items|length > {idx} %}}{{{{plan_items[{idx}].metric}}}}{{% endif %}}")
                set_run_font(run)

            print(f"  ✓ 項次 {idx + 1}")

    # ===== 第22行: 階段中目標 =====
    midterm_row = table.rows[22]
    if len(midterm_row.cells) > 1:
        cell = midterm_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        run = cell.paragraphs[0].add_run("{{midterm_goal}}")
        set_run_font(run)
        print("\n✓ 第22行: 階段中目標")

    # ===== 第23行: 階段末目標 =====
    final_row = table.rows[23]
    if len(final_row.cells) > 1:
        cell = final_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        run = cell.paragraphs[0].add_run("{{final_goal}}")
        set_run_font(run)
        print("✓ 第23行: 階段末目標")

    # ===== 第24行: 成果發表形式 =====
    presentation_row = table.rows[24]
    if len(presentation_row.cells) > 1:
        cell = presentation_row.cells[1]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        para = cell.paragraphs[0]
        text = (
            "{% if presentation_formats.get('靜態展') %}☑{% else %}□{% endif %}靜態展(PPT、書面報告、心得、自我省思…等)\n"
            "{% if presentation_formats.get('動態展') %}☑{% else %}□{% endif %}動態展(直播、影片撥放、實際展示、演出…等)\n"
            "{% if presentation_other %}☑{% else %}□{% endif %}其他 {{presentation_other}}"
        )
        run = para.add_run(text)
        set_run_font(run)
        print("✓ 第24行: 成果發表形式")

    # ===== 第25行: 手機使用規範 =====
    phone_row = table.rows[25]
    if len(phone_row.cells) > 8:
        cell = phone_row.cells[8]
        for para in cell.paragraphs:
            para.clear()
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        text = (
            "{% if phone_agreement == '同意' %}☑{% else %}□{% endif %}同意   "
            "{% if phone_agreement == '不同意' %}☑{% else %}□{% endif %}不同意\n"
            "(勾選不同意者本次申請將改列彈性學習)"
        )
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run)
        print("✓ 第25行: 手機使用規範")

    # 儲存檔案
    doc.save(output_path)
    print(f"\n{'=' * 80}")
    print(f"✅ 最終模板建立成功！")
    print(f"已儲存到: {output_path}")
    print(f"\n重點修正：")
    print(f"  1. ✅ 「是否繳交過」改為checkbox格式（是/否兩選項）")
    print(f"  2. ✅ 新增姓名欄位（第3-5行列6）")
    print(f"  3. ✅ 所有其他欄位保持完整")
    print(f"\n下一步: 複製到Docker容器")
    print(f'docker cp "附件一 復興自主學習申請表-新版.docx" self-learning-backend:/app/templates/application_template.docx')

except Exception as e:
    print(f"\n❌ 錯誤: {e}")
    import traceback
    traceback.print_exc()
