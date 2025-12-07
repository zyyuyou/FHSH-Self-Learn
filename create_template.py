#!/usr/bin/env python3
"""
將復興自主學習申請表.docx轉換為Jinja2模板
"""

from docx import Document
import shutil

# 讀取原始檔案
input_path = "./復興自主學習申請表.docx"
output_path = "./附件一 復興自主學習申請表-新版.docx"

# 先複製檔案
shutil.copy(input_path, output_path)

# 然後修改
doc = Document(output_path)

# 獲取主要表格
table = doc.tables[0] if doc.tables else None

if table:
    # 行2: 自主學習計畫名稱後面應該填入 {{title}}
    # 由於表格儲存格合併，我們需要找到正確的儲存格
    title_cell = table.rows[2].cells[1]  # 第2行第1列開始
    # 清空並設定變數
    for paragraph in title_cell.paragraphs:
        paragraph.text = ""
    title_cell.paragraphs[0].add_run("{{title}}")

    # 行3-6: 學生資料
    # 組長 (行4)
    member1_cells = table.rows[4].cells
    member1_cells[1].paragraphs[0].clear()
    member1_cells[1].paragraphs[0].add_run("{{member1.class_name}}")
    member1_cells[3].paragraphs[0].clear()
    member1_cells[3].paragraphs[0].add_run("{{member1.seat_number}}")
    member1_cells[4].paragraphs[0].clear()
    member1_cells[4].paragraphs[0].add_run("{{member1.student_id}}")
    # has_submitted 在行4列9
    submitted_cell = member1_cells[9]
    for para in submitted_cell.paragraphs:
        if "□是" in para.text:
            para.text = "{% if member1.has_submitted == '是' %}☑{% else %}□{% endif %}是  {% if member1.has_submitted == '第一次申請' %}☑{% else %}□{% endif %}第一次申請\n{% if member1.has_submitted == '否' %}☑{% else %}□{% endif %}否"

    # 組員2 (行5)
    member2_cells = table.rows[5].cells
    member2_cells[1].paragraphs[0].clear()
    member2_cells[1].paragraphs[0].add_run("{{member2.class_name}}")
    member2_cells[3].paragraphs[0].clear()
    member2_cells[3].paragraphs[0].add_run("{{member2.seat_number}}")
    member2_cells[4].paragraphs[0].clear()
    member2_cells[4].paragraphs[0].add_run("{{member2.student_id}}")

    # 組員3 (行6)
    member3_cells = table.rows[6].cells
    member3_cells[1].paragraphs[0].clear()
    member3_cells[1].paragraphs[0].add_run("{{member3.class_name}}")
    member3_cells[3].paragraphs[0].clear()
    member3_cells[3].paragraphs[0].add_run("{{member3.seat_number}}")
    member3_cells[4].paragraphs[0].clear()
    member3_cells[4].paragraphs[0].add_run("{{member3.student_id}}")

    # 行7: 學習動機
    motivation_cell = table.rows[7].cells[1]
    motivation_cell.paragraphs[0].clear()
    motivation_cell.paragraphs[0].add_run("{{motivation}}")

    # 行8: 學習類別
    category_cell = table.rows[8].cells[1]
    for para in category_cell.paragraphs:
        para.clear()
    # 新增學習類別的Jinja2模板
    para = category_cell.paragraphs[0]
    para.add_run("{% if learning_categories['閱讀計畫'] %}☑{% else %}□{% endif %}閱讀計畫 ")
    para.add_run("{% if learning_categories['志工服務'] %}☑{% else %}□{% endif %}志工服務 ")
    para.add_run("{% if learning_categories['專題研究'] %}☑{% else %}□{% endif %}專題研究 ")
    para.add_run("{% if learning_categories['藝文創作'] %}☑{% else %}□{% endif %}藝文創作 ")
    para.add_run("{% if learning_categories['技藝學習'] %}☑{% else %}□{% endif %}技藝學習 ")
    para.add_run("{% if learning_categories['競賽準備'] %}☑{% else %}□{% endif %}競賽準備\n")
    para.add_run("{% if learning_categories['實作體驗'] %}☑{% else %}□{% endif %}實作體驗 ")
    para.add_run("{% if learning_categories['課程延伸'] %}☑{% else %}□{% endif %}課程延伸 ")
    para.add_run("□其他：{{learning_category_other}}")

    # 行10: 學習裝置需求
    equipment_cell = table.rows[10].cells[1]
    equipment_cell.paragraphs[0].clear()
    equipment_cell.paragraphs[0].add_run("{{equipment_needs}}")

    # 行11: 學習方法(參考資料)
    reference_cell = table.rows[11].cells[1]
    reference_cell.paragraphs[0].clear()
    reference_cell.paragraphs[0].add_run("{% for ref in references %}{{ ref.book_title }} / {{ ref.author }} / {{ ref.publisher }}{% if ref.link %} / {{ ref.link }}{% endif %}\n{% endfor %}")

    # 行12: 預期成效
    outcome_cell = table.rows[12].cells[1]
    outcome_cell.paragraphs[0].clear()
    outcome_cell.paragraphs[0].add_run("{{expected_outcome}}")

    # 行14-22: 學習內容規劃 (9個項次)
    for i in range(9):
        row_idx = 14 + i
        if row_idx < len(table.rows):
            row = table.rows[row_idx]
            # 檢查是否有足夠的plan_items
            row.cells[2].paragraphs[0].clear()
            row.cells[2].paragraphs[0].add_run("{% if plan_items[" + str(i) + "] %}{{plan_items[" + str(i) + "].date}}{% endif %}")
            row.cells[4].paragraphs[0].clear()
            row.cells[4].paragraphs[0].add_run("{% if plan_items[" + str(i) + "] %}{{plan_items[" + str(i) + "].hours}}{% endif %}")
            row.cells[5].paragraphs[0].clear()
            row.cells[5].paragraphs[0].add_run("{% if plan_items[" + str(i) + "] %}{{plan_items[" + str(i) + "].content}}{% endif %}")
            row.cells[7].paragraphs[0].clear()
            row.cells[7].paragraphs[0].add_run("{% if plan_items[" + str(i) + "] %}{{plan_items[" + str(i) + "].metric}}{% endif %}")

    # 行23: 階段中目標
    midterm_cell = table.rows[23].cells[1]
    midterm_cell.paragraphs[0].clear()
    midterm_cell.paragraphs[0].add_run("{{midterm_goal}}")

    # 行24: 階段末目標
    final_cell = table.rows[24].cells[1]
    final_cell.paragraphs[0].clear()
    final_cell.paragraphs[0].add_run("{{final_goal}}")

    # 行25: 成果發表形式
    presentation_cell = table.rows[25].cells[1]
    for para in presentation_cell.paragraphs:
        para.clear()
    para = presentation_cell.paragraphs[0]
    para.add_run("{% if '靜態展' in presentation_format %}☑{% else %}□{% endif %}靜態展(PPT、書面報告、心得、自我省思…等) \n")
    para.add_run("{% if '動態展' in presentation_format %}☑{% else %}□{% endif %}動態展 (直播、影片撥放、實際展示、演出…等) \n")
    para.add_run("□其他{{presentation_other}}")

# 儲存模板
doc.save(output_path)
print(f"模板已建立: {output_path}")
