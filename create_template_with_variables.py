#!/usr/bin/env python3
"""
創建帶有 Jinja2 變量的 Word 模板
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

def create_template():
    """創建新的模板文件"""
    doc = Document()

    # 設置頁面
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4 height
    section.page_width = Inches(8.27)    # A4 width
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # 標題
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('臺北市立復興高級中學自主學習計畫 學生申請表')
    run.font.size = Pt(18)
    run.bold = True

    doc.add_paragraph()

    # 創建表格
    # 基本資訊表格
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'

    # 計畫名稱
    table.rows[0].cells[0].text = '計畫名稱'
    table.rows[0].cells[1].text = '{{title}}'

    # 申請期間
    table.rows[1].cells[0].text = '申請期間'
    table.rows[1].cells[1].text = '{{apply_date_start}} 至 {{apply_date_end}}'

    doc.add_paragraph()

    # 組員資訊表格
    members_title = doc.add_paragraph('一、組員資訊')
    members_title.runs[0].bold = True

    members_table = doc.add_table(rows=4, cols=6)
    members_table.style = 'Table Grid'

    # 表頭
    headers = ['序號', '學號', '班級', '座號', '姓名', '是否繳交過自主學習成果']
    for i, header in enumerate(headers):
        members_table.rows[0].cells[i].text = header

    # 組員資料 (使用 Jinja2 循環)
    for i in range(1, 4):
        members_table.rows[i].cells[0].text = f'{{% if members[{i-1}] %}}{i}{{% endif %}}'
        members_table.rows[i].cells[1].text = f'{{% if members[{i-1}] %}}{{{{members[{i-1}].student_id}}}}{{% endif %}}'
        members_table.rows[i].cells[2].text = f'{{% if members[{i-1}] %}}{{{{members[{i-1}].class_name}}}}{{% endif %}}'
        members_table.rows[i].cells[3].text = f'{{% if members[{i-1}] %}}{{{{members[{i-1}].seat_number}}}}{{% endif %}}'
        members_table.rows[i].cells[4].text = f'{{% if members[{i-1}] %}}{{{{members[{i-1}].name}}}}{{% endif %}}'
        members_table.rows[i].cells[5].text = f'{{% if members[{i-1}] %}}{{{{members[{i-1}].has_submitted}}}}{{% endif %}}'

    doc.add_paragraph()

    # 學習動機
    motivation_title = doc.add_paragraph('二、學習動機')
    motivation_title.runs[0].bold = True
    doc.add_paragraph('{{motivation}}')

    doc.add_paragraph()

    # 學習類別
    category_title = doc.add_paragraph('三、學習類別')
    category_title.runs[0].bold = True
    doc.add_paragraph('{{learning_categories}}')

    doc.add_paragraph()

    # 學習環境需求
    env_title = doc.add_paragraph('四、學習環境需求')
    env_title.runs[0].bold = True
    doc.add_paragraph('{{env_needs}}')

    doc.add_paragraph()

    # 學習內容規劃
    plan_title = doc.add_paragraph('五、學習內容規劃')
    plan_title.runs[0].bold = True

    plan_table = doc.add_table(rows=1, cols=5)
    plan_table.style = 'Table Grid'

    # 表頭
    plan_headers = ['序號', '日期', '學習內容', '時數', '學生自訂檢核指標']
    for i, header in enumerate(plan_headers):
        plan_table.rows[0].cells[i].text = header

    # 計畫項目 (使用固定行數 + 條件渲染)
    # 創建 10 行，使用條件來顯示有數據的行
    for i in range(10):
        row = plan_table.add_row()
        # Use string concatenation instead of f-strings to avoid escaping issues
        row.cells[0].text = '{% if plan_items|length > ' + str(i) + ' %}{{ plan_items[' + str(i) + '].number }}{% endif %}'
        row.cells[1].text = '{% if plan_items|length > ' + str(i) + ' %}{{ plan_items[' + str(i) + '].date }}{% endif %}'
        row.cells[2].text = '{% if plan_items|length > ' + str(i) + ' %}{{ plan_items[' + str(i) + '].content }}{% endif %}'
        row.cells[3].text = '{% if plan_items|length > ' + str(i) + ' %}{{ plan_items[' + str(i) + '].hours }}{% endif %}'
        row.cells[4].text = '{% if plan_items|length > ' + str(i) + ' %}{{ plan_items[' + str(i) + '].metric }}{% endif %}'

    doc.add_paragraph()

    # 成果發表形式
    presentation_title = doc.add_paragraph('六、成果發表形式')
    presentation_title.runs[0].bold = True
    doc.add_paragraph('{{presentation_format}}')
    if_other = doc.add_paragraph('{% if presentation_other %}其他：{{presentation_other}}{% endif %}')

    doc.add_paragraph()
    doc.add_paragraph()

    # 簽名欄
    signature_title = doc.add_paragraph('七、簽名')
    signature_title.runs[0].bold = True

    signature_table = doc.add_table(rows=2, cols=4)
    signature_table.style = 'Table Grid'

    signature_table.rows[0].cells[0].text = '學生簽名'
    signature_table.rows[0].cells[1].text = '家長簽名'
    signature_table.rows[0].cells[2].text = '導師簽名'
    signature_table.rows[0].cells[3].text = '輔導老師簽名'

    signature_table.rows[1].cells[0].text = '{{student_signature}}'
    signature_table.rows[1].cells[1].text = '{{parent_signature}}'
    signature_table.rows[1].cells[2].text = '{{homeroom_teacher_signature}}'
    signature_table.rows[1].cells[3].text = '{{counselor_signature}}'

    doc.add_paragraph()
    doc.add_paragraph()

    # 審核結果
    review_title = doc.add_paragraph('八、審核結果')
    review_title.runs[0].bold = True

    review_table = doc.add_table(rows=2, cols=2)
    review_table.style = 'Table Grid'

    review_table.rows[0].cells[0].text = '審核狀態'
    review_table.rows[0].cells[1].text = '{{status}}'

    review_table.rows[1].cells[0].text = '教師評語'
    review_table.rows[1].cells[1].text = '{{comment}}'

    # 保存
    output_path = 'application_template_with_variables.docx'
    doc.save(output_path)
    print(f'✓ 模板已創建: {output_path}')
    return output_path

if __name__ == '__main__':
    create_template()
