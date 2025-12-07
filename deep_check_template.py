#!/usr/bin/env python3
from docx import Document

template_path = "./附件一 復興自主學習申請表-新版.docx"
doc = Document(template_path)
table = doc.tables[0]

print("深入檢查第3行第9列的結構:")
print("=" * 80)

row = table.rows[3]
cell = row.cells[9]

print(f"單元格文字: {repr(cell.text)}")
print(f"\n段落數量: {len(cell.paragraphs)}")

for i, para in enumerate(cell.paragraphs):
    print(f"\n段落 {i}:")
    print(f"  段落文字: {repr(para.text)}")
    print(f"  run數量: {len(para.runs)}")

    for j, run in enumerate(para.runs):
        print(f"  run {j}: {repr(run.text)}")
