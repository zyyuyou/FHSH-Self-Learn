#!/usr/bin/env python3
"""
修復合併儲存格問題 - 取消不必要的合併
"""
from docx import Document
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import parse_xml
import shutil

input_path = "./復興自主學習申請表.docx"
output_path = "./復興自主學習申請表-已取消合併.docx"

# 備份
backup_path = "./復興自主學習申請表.docx.bak_merge"
shutil.copy(input_path, backup_path)
print(f"✅ 已備份至: {backup_path}")

print("\n" + "=" * 80)
print("修復合併儲存格問題...")
print("=" * 80)

doc = Document(input_path)
table = doc.tables[0]

# 需要處理的行
rows_to_fix = [
    (6, "學習動機"),
    (7, "學習類別"),
    (8, "學習環境需求"),
    (9, "學習裝置需求"),
    (11, "預期成效"),
    (22, "階段中目標"),
    (23, "階段末目標"),
    (24, "成果發表形式"),
]

for row_idx, row_name in rows_to_fix:
    print(f"\n處理第{row_idx}行 ({row_name})...")

    if row_idx >= len(table.rows):
        print(f"  ⚠️  跳過（行不存在）")
        continue

    row = table.rows[row_idx]

    # 檢查合併情況
    if len(row.cells) > 1:
        # 獲取第二個cell（索引1）的底層元素
        first_content_cell = row.cells[1]
        tc_element = first_content_cell._tc

        # 檢查是否有 gridSpan（橫向合併）
        tcPr = tc_element.find(qn('w:tcPr'))
        if tcPr is not None:
            gridSpan = tcPr.find(qn('w:gridSpan'))
            if gridSpan is not None:
                span_value = gridSpan.get(qn('w:val'))
                print(f"  發現合併: gridSpan = {span_value}")

                # 移除 gridSpan（取消合併）
                tcPr.remove(gridSpan)
                print(f"  ✅ 已取消合併")
            else:
                print(f"  ℹ️  沒有gridSpan（可能使用其他合併方式）")
        else:
            print(f"  ℹ️  沒有tcPr元素")

# 儲存
doc.save(output_path)
print("\n" + "=" * 80)
print(f"✅ 已儲存到: {output_path}")
print("=" * 80)

# 驗證
print("\n驗證修復結果...")
doc_verify = Document(output_path)
table_verify = doc_verify.tables[0]

for row_idx, row_name in rows_to_fix:
    if row_idx < len(table_verify.rows):
        row = table_verify.rows[row_idx]
        if len(row.cells) > 1:
            tc = row.cells[1]._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                gridSpan = tcPr.find(qn('w:gridSpan'))
                if gridSpan is not None:
                    print(f"  ⚠️  第{row_idx}行仍有gridSpan")
                else:
                    print(f"  ✅ 第{row_idx}行已無gridSpan")
            else:
                print(f"  ✅ 第{row_idx}行已無tcPr")

print("\n" + "=" * 80)
print("修復完成！下一步：使用此檔案重新生成模板")
print("=" * 80)
