#!/usr/bin/env python3
"""
正確修復Word模板中的換行符和空段落
"""
from docx import Document
from docx.oxml import OxmlElement
import shutil

input_path = "./附件一 復興自主學習申請表-新版.docx"
output_path = "./附件一 復興自主學習申請表-新版-已修復.docx"

# 先備份
backup_path = "./附件一 復興自主學習申請表-新版.docx.bak4"
shutil.copy(input_path, backup_path)
print(f"✅ 已備份至: {backup_path}")

print("\n" + "=" * 80)
print("修復模板中的換行符問題...")
print("=" * 80)

doc = Document(input_path)
table = doc.tables[0]

fixed_runs = 0
removed_paragraphs = 0

# 遍歷所有單元格
for row_idx, row in enumerate(table.rows):
    for cell_idx, cell in enumerate(row.cells):
        # 第一步：修復run中的換行符
        for para in cell.paragraphs:
            for run in para.runs:
                original_text = run.text

                # 檢查是否包含換行符
                if '\n' in original_text:
                    # 對於「是否繳交過」欄位，保留「是」和「否」之間的換行，但移除結尾換行
                    if '{% if' in original_text and '是' in original_text and '否' in original_text:
                        # 這是「是否繳交過」欄位
                        # 保留一個中間換行，移除結尾換行
                        parts = original_text.split('\n')
                        # 過濾掉空字串
                        non_empty_parts = [p for p in parts if p.strip()]
                        if len(non_empty_parts) >= 2:
                            # 用一個換行連線非空部分
                            cleaned_text = '\n'.join(non_empty_parts)
                        else:
                            cleaned_text = '\n'.join(parts[:-1]) if parts[-1] == '' else original_text.rstrip('\n')
                    else:
                        # 其他欄位：移除所有換行符
                        cleaned_text = original_text.replace('\n', '')

                    if cleaned_text != original_text:
                        run.text = cleaned_text
                        fixed_runs += 1

                        if fixed_runs <= 5:
                            print(f"\n修復run #{fixed_runs}: 第{row_idx}行 第{cell_idx}列")
                            print(f"  原: {repr(original_text[:60])}")
                            print(f"  新: {repr(cleaned_text[:60])}")

        # 第二步：移除多餘的空段落（保留至少一個段落）
        if len(cell.paragraphs) > 1:
            # 檢查並移除空段落
            paragraphs_to_remove = []

            for i, para in enumerate(cell.paragraphs):
                # 如果是最後一個段落之後的空段落，標記刪除
                if i > 0 and para.text.strip() == '' and len(para.runs) == 0:
                    paragraphs_to_remove.append(para)

            # 刪除空段落
            for para in paragraphs_to_remove:
                p_element = para._element
                p_element.getparent().remove(p_element)
                removed_paragraphs += 1

print(f"\n{'=' * 80}")
print(f"✅ 修復了 {fixed_runs} 個run的換行符")
print(f"✅ 移除了 {removed_paragraphs} 個空段落")
print(f"{'=' * 80}")

# 儲存
doc.save(output_path)
print(f"\n✅ 已儲存到: {output_path}")

# 驗證
print("\n" + "=" * 80)
print("驗證修復結果...")
print("=" * 80)

doc_verify = Document(output_path)
table_verify = doc_verify.tables[0]

# 檢查「是否繳交過」欄位
for student_idx, row_idx in enumerate([3, 4, 5], 1):
    row = table_verify.rows[row_idx]
    cell = row.cells[9]

    print(f"\n學生{student_idx} (第{row_idx}行):")
    print(f"  單元格文字: {repr(cell.text)}")
    print(f"  段落數: {len(cell.paragraphs)}")

    # 檢查是否還有結尾換行
    if cell.text.endswith('\n'):
        print(f"  ⚠️  仍有結尾換行")
    else:
        print(f"  ✅ 無結尾換行")

    # 檢查段落
    for i, para in enumerate(cell.paragraphs):
        if para.text.strip():
            print(f"  段落{i}: {repr(para.text[:60])}")

print("\n" + "=" * 80)
print("修復完成！")
print("=" * 80)
