#!/usr/bin/env python3
"""
修復Word模板中的多餘換行符
"""
from docx import Document
import shutil

input_path = "./附件一 復興自主學習申請表-新版.docx"
output_path = "./附件一 復興自主學習申請表-新版-修復換行.docx"

# 先備份
backup_path = "./附件一 復興自主學習申請表-新版.docx.bak3"
shutil.copy(input_path, backup_path)
print(f"✅ 已備份至: {backup_path}")

print("\n" + "=" * 80)
print("修復模板中的多餘換行符...")
print("=" * 80)

doc = Document(input_path)
table = doc.tables[0]

fixed_count = 0

# 遍歷所有單元格
for row_idx, row in enumerate(table.rows):
    for cell_idx, cell in enumerate(row.cells):
        # 遍歷單元格中的所有段落
        for para_idx, para in enumerate(cell.paragraphs):
            # 遍歷段落中的所有 run
            for run_idx, run in enumerate(para.runs):
                original_text = run.text

                # 檢查是否包含 Jinja2 語法且結尾有換行
                if ('{%' in original_text or '{{' in original_text) and original_text.endswith('\n'):
                    # 移除結尾的所有換行符（可能有多個）
                    cleaned_text = original_text.rstrip('\n')

                    # 如果文字被修改了
                    if cleaned_text != original_text:
                        run.text = cleaned_text
                        fixed_count += 1

                        # 顯示修復資訊（只顯示前幾個）
                        if fixed_count <= 10:
                            original_display = original_text if len(original_text) < 50 else original_text[:47] + "..."
                            cleaned_display = cleaned_text if len(cleaned_text) < 50 else cleaned_text[:47] + "..."
                            print(f"\n修復 #{fixed_count}: 第{row_idx}行 第{cell_idx}列")
                            print(f"  原始: {repr(original_display)}")
                            print(f"  修復: {repr(cleaned_display)}")
                        elif fixed_count == 11:
                            print(f"\n... (還有更多修復，不再顯示詳情)")

print("\n" + "=" * 80)
print(f"✅ 總共修復了 {fixed_count} 個換行符問題")
print("=" * 80)

# 儲存修復後的檔案
doc.save(output_path)
print(f"\n✅ 已儲存到: {output_path}")

# 驗證修復
print("\n" + "=" * 80)
print("驗證修復結果...")
print("=" * 80)

doc_verify = Document(output_path)
table_verify = doc_verify.tables[0]

remaining_issues = 0
for row in table_verify.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                text = run.text
                if ('{%' in text or '{{' in text) and text.endswith('\n'):
                    remaining_issues += 1

if remaining_issues == 0:
    print("✅ 驗證透過！沒有發現剩餘的換行符問題")
else:
    print(f"⚠️  仍有 {remaining_issues} 個換行符問題")

# 特別檢查「是否繳交過」欄位
print("\n" + "=" * 80)
print("特別驗證「是否繳交過」欄位...")
print("=" * 80)

for row_idx in [3, 4, 5]:
    row = table_verify.rows[row_idx]
    cell = row.cells[9]
    text = cell.text
    print(f"\n第{row_idx}行:")
    print(f"  {repr(text)}")
    if text.endswith('\n'):
        print(f"  ⚠️  仍有結尾換行符")
    else:
        print(f"  ✅ 無結尾換行符")

print("\n" + "=" * 80)
print(f"修復完成！")
print(f"請將修復後的檔案複製到 Docker:")
print(f'docker cp "{output_path}" self-learning-backend:/app/templates/application_template.docx')
print("=" * 80)
