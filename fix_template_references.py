#!/usr/bin/env python3
"""
修正模板中參考資料的內嵌表格
"""

from docx import Document
import shutil

# 路徑設定
input_path = "./復興自主學習申請表.docx"
output_path = "./附件一 復興自主學習申請表-新版.docx"

# 先複製原始檔案
shutil.copy(input_path, output_path)
print(f"已複製: {input_path} -> {output_path}")

# 開啟檔案
doc = Document(output_path)
table = doc.tables[0]

print(f"\n開始修正參考資料部分...")
print("=" * 80)

try:
    # 第10行包含參考資料的內嵌表格
    ref_row = table.rows[10]
    ref_cell = ref_row.cells[1]

    if len(ref_cell.tables) > 0:
        inner_table = ref_cell.tables[0]
        print(f"找到內嵌表格，共 {len(inner_table.rows)} 行")

        # 內嵌表格第0行是標題：書名 | 作者 | 出版社
        # 內嵌表格第1行是資料填入區

        # 我們需要為每個參考資料建立一行
        # 由於Jinja2在Word中處理迴圈比較複雜，我們先提供3行的模板

        # 先刪除現有的第1行（空白行）
        if len(inner_table.rows) > 1:
            # 注意：python-docx 無法直接刪除行，我們需要手動操作
            # 改為直接在第1行新增Jinja2迴圈

            data_row = inner_table.rows[1]

            # 書名（列0）
            cell_book = data_row.cells[0]
            for para in cell_book.paragraphs:
                para.clear()
            if len(cell_book.paragraphs) == 0:
                cell_book.add_paragraph()
            # 使用迴圈語法
            cell_book.paragraphs[0].add_run("{% for ref in references %}{{ref.book_title}}{% if ref.link %}\n連結: {{ref.link}}{% endif %}{% if not loop.last %}\n\n{% endif %}{% endfor %}")

            # 作者（列1）
            cell_author = data_row.cells[1]
            for para in cell_author.paragraphs:
                para.clear()
            if len(cell_author.paragraphs) == 0:
                cell_author.add_paragraph()
            cell_author.paragraphs[0].add_run("{% for ref in references %}{{ref.author}}{% if not loop.last %}\n\n{% endif %}{% endfor %}")

            # 出版社（列2）
            cell_publisher = data_row.cells[2]
            for para in cell_publisher.paragraphs:
                para.clear()
            if len(cell_publisher.paragraphs) == 0:
                cell_publisher.add_paragraph()
            cell_publisher.paragraphs[0].add_run("{% for ref in references %}{{ref.publisher}}{% if not loop.last %}\n\n{% endif %}{% endfor %}")

            print("✓ 已在內嵌表格中新增參考資料變數")
            print("  - 書名欄: {% for ref in references %}{{ref.book_title}}...")
            print("  - 作者欄: {% for ref in references %}{{ref.author}}...")
            print("  - 出版社欄: {% for ref in references %}{{ref.publisher}}...")
    else:
        print("❌ 未找到內嵌表格")

    # 儲存檔案
    doc.save(output_path)
    print(f"\n✅ 成功! 已儲存到: {output_path}")

except Exception as e:
    print(f"\n❌ 錯誤: {e}")
    import traceback
    traceback.print_exc()
