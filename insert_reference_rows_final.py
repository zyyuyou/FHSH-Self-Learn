#!/usr/bin/env python3
"""
正確插入參考資料行 - 複製學習內容規劃的行結構
"""
from docx import Document
from copy import deepcopy

template_path = "./附件一 復興自主學習申請表-新版.docx"

print(f"讀取模板: {template_path}")
doc = Document(template_path)

table = doc.tables[0]

# 策略：複製學習內容規劃的一個資料行（比如Row 13），
# 然後修改其內容為參考資料，並插入到Row 10之後

print("\n=== 找到學習內容規劃的行作為模板 ===")
# Row 13 是學習內容規劃的第一個資料行
template_row = table.rows[13]
print(f"使用 Row 13 作為模板")

# 先更新Row 10為參考資料的標題行
print("\n=== 更新 Row 10 為參考資料標題行 ===")
row_10 = table.rows[10]
# 保持 Cell 0 不變（學習方法(參考資料)）
# 更新其他cells
row_10.cells[1].text = "書名"
row_10.cells[2].text = ""
row_10.cells[3].text = ""
row_10.cells[4].text = "作者"
row_10.cells[5].text = "出版社"
row_10.cells[6].text = ""
row_10.cells[7].text = "連結"
row_10.cells[8].text = ""
row_10.cells[9].text = ""

print("已更新標題行")

# 獲取表格的底層XML
tbl = table._tbl
row_11_xml = table.rows[11]._element

# 插入3個參考資料行
print("\n=== 插入3個參考資料資料行 ===")
for idx in range(3):
    print(f"插入參考資料 {idx+1}...")

    # 深度複製Row 13的XML
    new_row_xml = deepcopy(template_row._element)

    # 獲取新行中的所有cells
    from docx.oxml.ns import qn
    cells_xml = new_row_xml.findall(qn('w:tc'))

    def clear_and_set_cell(cell_xml, text):
        """清空cell並設定新文字"""
        # 清除所有段落
        for p in list(cell_xml.findall(qn('w:p'))):
            cell_xml.remove(p)
        # 建立新段落
        new_p = cell_xml.makeelement(qn('w:p'))
        new_r = new_p.makeelement(qn('w:r'))
        new_t = new_r.makeelement(qn('w:t'))
        new_t.text = text
        new_r.append(new_t)
        new_p.append(new_r)
        cell_xml.append(new_p)

    # XML cells 結構：[0]=左邊合併, [1]=項次, [2]=日期, [3]=時數, [4]=內容, [5]=檢核指標
    # 我們要改成：[0]=左邊合併, [1]=空, [2]=書名, [3]=作者, [4]=出版社, [5]=連結

    # Cell 0: 空（與標題合併）
    clear_and_set_cell(cells_xml[0], "")

    # Cell 1: 空（原本是項次）
    clear_and_set_cell(cells_xml[1], "")

    # Cell 2: 書名（原本是日期）
    book_var = f"{{% if references|length > {idx} %}}{{{{references[{idx}].book_title}}}}{{% endif %}}"
    clear_and_set_cell(cells_xml[2], book_var)

    # Cell 3: 作者（原本是時數）
    author_var = f"{{% if references|length > {idx} %}}{{{{references[{idx}].author}}}}{{% endif %}}"
    clear_and_set_cell(cells_xml[3], author_var)

    # Cell 4: 出版社（原本是內容）
    pub_var = f"{{% if references|length > {idx} %}}{{{{references[{idx}].publisher}}}}{{% endif %}}"
    clear_and_set_cell(cells_xml[4], pub_var)

    # Cell 5: 連結（原本是檢核指標）
    link_var = f"{{% if references|length > {idx} %}}{{{{references[{idx}].link}}}}{{% endif %}}"
    clear_and_set_cell(cells_xml[5], link_var)

    # 在Row 11之前插入新行
    row_11_xml.addprevious(new_row_xml)

print("\n✅ 完成！")

doc.save(template_path)
print(f"\n已儲存: {template_path}")

# 驗證
print("\n=== 驗證結果 ===")
doc2 = Document(template_path)
table2 = doc2.tables[0]
for i in range(10, 15):
    if i >= len(table2.rows):
        break
    row = table2.rows[i]
    cell_0 = row.cells[0].text[:30].replace('\n', ' ')
    cell_2 = row.cells[2].text[:50].replace('\n', ' ')
    print(f"Row {i}: [{cell_0}] | [{cell_2}]")
