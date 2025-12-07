#!/usr/bin/env python3
"""
測試渲染並檢查 Word 檔案內容
"""
from docxtpl import DocxTemplate
from jinja2 import Environment, Undefined
from docx import Document

class SilentUndefined(Undefined):
    """靜默處理未定義變數"""
    def _fail_with_undefined_error(self, *args, **kwargs):
        return ''
    __add__ = __radd__ = __mul__ = __rmul__ = __div__ = __rdiv__ = \
    __truediv__ = __rtruediv__ = __floordiv__ = __rfloordiv__ = \
    __mod__ = __rmod__ = __pos__ = __neg__ = \
    lambda self, other: self._fail_with_undefined_error()
    __lt__ = __le__ = __gt__ = __ge__ = __eq__ = __ne__ = \
    __hash__ = lambda self, other: self._fail_with_undefined_error()
    __getitem__ = lambda self, other: self._fail_with_undefined_error()
    def __str__(self):
        return ''
    def __len__(self):
        return 0
    def __iter__(self):
        return iter([])
    def __bool__(self):
        return False

template_path = "./附件一 復興自主學習申請表-新版.docx"
output_path = "./debug_rendered.docx"

# 測試資料：只選閱讀計畫
test_data = {
    'title': '測試',
    'apply_date_start': '2025-01-01',
    'apply_date_end': '2025-06-30',

    'member1': {
        'student_id': '11430001',
        'class_name': '101',
        'seat_number': '1',
        'name': '測試學生1',
        'has_submitted': '否'
    },
    'member2': {
        'student_id': '',
        'class_name': '',
        'seat_number': '',
        'name': '',
        'has_submitted': ''
    },
    'member3': {
        'student_id': '',
        'class_name': '',
        'seat_number': '',
        'name': '',
        'has_submitted': ''
    },

    'motivation': '測試動機',

    # 學習類別：只選「閱讀計畫」
    'learning_categories': {
        '閱讀計畫': True,
        '志工服務': False,
        '專題研究': False,
        '藝文創作': False,
        '技藝學習': False,
        '競賽準備': False,
        '實作體驗': False,
        '課程延伸': False,
    },
    'learning_category_other': '',  # 空的「其他」

    'references': [],
    'expected_outcome': '測試成效',
    'equipment_needs': '電腦',

    'env_needs': {
        '自習室': True,
        '數位閱讀室': False,
        '雲端教室': False,
        '美力教室': False,
    },
    'env_other': '',

    'plan_items': [
        {
            'date': '1/15',
            'content': '測試內容',
            'hours': '2',
            'metric': '測試指標'
        }
    ],

    'midterm_goal': '中期目標',
    'final_goal': '最終目標',

    # 成果發表：只選「靜態展」
    'presentation_formats': {
        '靜態展': True,
        '動態展': False,
    },
    'presentation_other': '',  # 空的「其他」

    'phone_agreement': '同意',
    'status': '審核中',
    'comment': '',
}

print("=" * 80)
print("渲染模板...")
print("=" * 80)

doc = DocxTemplate(template_path)
jinja_env = Environment(undefined=SilentUndefined)
doc.render(test_data, jinja_env=jinja_env)
doc.save(output_path)

print(f"\n✅ 已生成: {output_path}")

# 檢查渲染後的內容
print("\n" + "=" * 80)
print("檢查渲染結果...")
print("=" * 80)

doc_result = Document(output_path)
table = doc_result.tables[0]

print("\n學習類別（第7行）:")
row = table.rows[7]
print(f"Cells 數: {len(row.cells)}")

for i in range(min(3, len(row.cells))):
    cell = row.cells[i]
    print(f"\nCell {i}:")
    print(f"  文字長度: {len(cell.text)}")
    print(f"  段落數: {len(cell.paragraphs)}")

    for j, para in enumerate(cell.paragraphs):
        if para.text:  # 只顯示非空段落
            print(f"  段落 {j} (長度 {len(para.text)}): {repr(para.text[:100])}")
            print(f"    runs 數: {len(para.runs)}")
            for k, run in enumerate(para.runs):
                if run.text:
                    print(f"      run {k}: {repr(run.text[:50])}")

print("\n" + "=" * 80)
print("\n成果發表形式（第24行）:")
row = table.rows[24]
print(f"Cells 數: {len(row.cells)}")

for i in range(min(3, len(row.cells))):
    cell = row.cells[i]
    print(f"\nCell {i}:")
    print(f"  文字長度: {len(cell.text)}")
    print(f"  段落數: {len(cell.paragraphs)}")

    for j, para in enumerate(cell.paragraphs):
        if para.text:  # 只顯示非空段落
            print(f"  段落 {j} (長度 {len(para.text)}): {repr(para.text[:100])}")
            print(f"    runs 數: {len(para.runs)}")
            for k, run in enumerate(para.runs):
                if run.text:
                    print(f"      run {k}: {repr(run.text[:50])}")
