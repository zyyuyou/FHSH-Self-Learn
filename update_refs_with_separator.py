#!/usr/bin/env python3
"""
更新參考資料欄位 - 在同一個cell中用清晰的分隔線分開每個參考資料
"""
from docx import Document

template_path = "./附件一 復興自主學習申請表-新版.docx"

print(f"讀取模板: {template_path}")
doc = Document(template_path)

table = doc.tables[0]
row_10 = table.rows[10]

print("\n更新第10行（學習方法(參考資料)）...")

# 獲取Cell 1並清空
cell_1 = row_10.cells[1]
cell_1.text = ""

# 新增新段落，使用表格樣式來分隔每個參考資料
para = cell_1.paragraphs[0]

# 建立一個小表格式的顯示，每個參考資料包含：
# 序號、書名、作者、出版社、連結（分多行）
template_text = """{% for ref in references %}【參考資料 {{ loop.index }}】
書名：{{ ref.book_title }}
作者：{{ ref.author }}
出版社：{{ ref.publisher }}{% if ref.link %}
連結：{{ ref.link }}{% endif %}
{% if not loop.last %}
----------------------------
{% endif %}{% endfor %}"""

para.text = template_text

print("  已更新參考資料模板")
print("  每個參考資料包含：標題、書名、作者、出版社、連結")
print("  參考資料之間用分隔線區分")

doc.save(template_path)
print(f"\n✅ 已儲存: {template_path}")
