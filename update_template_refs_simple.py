#!/usr/bin/env python3
"""
簡單更新參考資料欄位 - 用段落分隔多個參考資料
"""
from docx import Document

template_path = "./附件一 復興自主學習申請表-新版.docx"

print(f"讀取模板: {template_path}")
doc = Document(template_path)

table = doc.tables[0]
row_10 = table.rows[10]

print("\n更新第10行（學習方法(參考資料)）...")

# 獲取Cell 1並清空
cell_1 = row_10.cells[1]
# 清空現有內容
cell_1.text = ""

# 新增新段落作為參考資料模板
# 每個參考資料一行，用Jinja2迴圈
para = cell_1.paragraphs[0]  # 使用第一個段落
para.text = """{% for ref in references %}{{ loop.index }}. {{ ref.book_title }} / {{ ref.author }} / {{ ref.publisher }}{% if ref.link %} / {{ ref.link }}{% endif %}
{% endfor %}"""

print("  已更新參考資料模板")
print(f"  使用Jinja2迴圈，每個參考資料獨立一行")

doc.save(template_path)
print(f"\n✅ 已儲存: {template_path}")
