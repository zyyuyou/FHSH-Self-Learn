#!/usr/bin/env python3
"""
更新參考資料欄位 - 新增多個表格行，類似學習內容規劃
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy

template_path = "./附件一 復興自主學習申請表-新版.docx"

print(f"讀取模板: {template_path}")
doc = Document(template_path)

table = doc.tables[0]

print("\n=== 當前第10行結構 ===")
row_10 = table.rows[10]
print(f"Row 10: {len(row_10.cells)} cells")
for i in range(min(5, len(row_10.cells))):
    print(f"  Cell {i}: {repr(row_10.cells[i].text[:50])}")

# 我們需要在第10行（標題行）之後插入新的行來顯示參考資料
# 參考學習內容規劃的結構：標題行 + 資料行

# 第10行是標題行，需要將它改為類似「學習內容規劃」的格式
# 標題行結構：Cell 0: 標題（跨多行）, Cell 1: 書名, Cell 2-3: 作者, Cell 4-5: 出版社, Cell 6-9: 連結

print("\n=== 更新第10行為標題行 ===")
# 清空並設定標題行的各個欄位
cells = row_10.cells

# Cell 0: 保持標題不變
# Cell 1: 書名
cells[1].text = ""
cells[1].paragraphs[0].text = "書名"

# Cell 2-3: 作者（合併顯示）
cells[2].text = ""
cells[2].paragraphs[0].text = "作者"

# Cell 4-5: 出版社
cells[4].text = ""
cells[4].paragraphs[0].text = "出版社"

# Cell 6-9: 連結
cells[6].text = ""
cells[6].paragraphs[0].text = "連結（電子書籍或電子期刊）"

print("  已設定標題行欄位")

# 現在在第10行之後插入3個資料行（支援3個參考資料）
print("\n=== 插入參考資料資料行 ===")

# 獲取第11行的XML元素，我們將在它之前插入新行
row_11 = table.rows[11]
row_11_xml = row_11._element

# 建立3個新行
for ref_idx in range(3):
    print(f"  插入參考資料 {ref_idx + 1} 的行...")

    # 複製第10行的結構來建立新行
    new_row = deepcopy(row_10._element)

    # 獲取新行的所有cells
    new_cells = new_row.findall(qn('w:tc'))

    # 更新Cell 0（序號，與第10行合併）
    # Cell 0 應該保持與標題行的合併狀態

    # Cell 1: 書名變數
    if len(new_cells) > 1:
        cell_1 = new_cells[1]
        # 清空cell內容
        for p in cell_1.findall(qn('w:p')):
            cell_1.remove(p)
        # 新增新段落
        new_p = OxmlElement('w:p')
        new_r = OxmlElement('w:r')
        new_t = OxmlElement('w:t')
        new_t.text = f"{{% if references|length > {ref_idx} %}}{{{{references[{ref_idx}].book_title}}}}{{% endif %}}"
        new_r.append(new_t)
        new_p.append(new_r)
        cell_1.append(new_p)

    # Cell 2-3: 作者變數
    if len(new_cells) > 2:
        cell_2 = new_cells[2]
        for p in cell_2.findall(qn('w:p')):
            cell_2.remove(p)
        new_p = OxmlElement('w:p')
        new_r = OxmlElement('w:r')
        new_t = OxmlElement('w:t')
        new_t.text = f"{{% if references|length > {ref_idx} %}}{{{{references[{ref_idx}].author}}}}{{% endif %}}"
        new_r.append(new_t)
        new_p.append(new_r)
        cell_2.append(new_p)

    # Cell 4-5: 出版社變數
    if len(new_cells) > 4:
        cell_4 = new_cells[4]
        for p in cell_4.findall(qn('w:p')):
            cell_4.remove(p)
        new_p = OxmlElement('w:p')
        new_r = OxmlElement('w:r')
        new_t = OxmlElement('w:t')
        new_t.text = f"{{% if references|length > {ref_idx} %}}{{{{references[{ref_idx}].publisher}}}}{{% endif %}}"
        new_r.append(new_t)
        new_p.append(new_r)
        cell_4.append(new_p)

    # Cell 6-9: 連結變數
    if len(new_cells) > 6:
        cell_6 = new_cells[6]
        for p in cell_6.findall(qn('w:p')):
            cell_6.remove(p)
        new_p = OxmlElement('w:p')
        new_r = OxmlElement('w:r')
        new_t = OxmlElement('w:t')
        new_t.text = f"{{% if references|length > {ref_idx} %}}{{{{references[{ref_idx}].link}}}}{{% endif %}}"
        new_r.append(new_t)
        new_p.append(new_r)
        cell_6.append(new_p)

    # 在第11行之前插入新行
    row_11_xml.addprevious(new_row)

print("\n✅ 完成！已插入3個參考資料資料行")

doc.save(template_path)
print(f"\n已儲存: {template_path}")
