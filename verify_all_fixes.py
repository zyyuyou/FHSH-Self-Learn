#!/usr/bin/env python3
"""
驗證所有修正是否正確套用到最終模板
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

template_path = "./附件一 復興自主學習申請表-新版.docx"

print("=" * 80)
print("驗證所有修正...")
print("=" * 80)

doc = Document(template_path)

# 檢查1: Google表單文字是否已刪除
print("\n【檢查1】Google表單提示文字")
if len(doc.paragraphs) > 0:
    first_para_text = doc.paragraphs[0].text.strip()
    if 'Google' in first_para_text:
        print(f"  ❌ 仍存在: {first_para_text}")
    else:
        print(f"  ✅ 已刪除 (第一段落內容: '{first_para_text}')")
else:
    print("  ✅ 已刪除")

table = doc.tables[0]

# 檢查2: 是否繳交過 - checkbox格式
print("\n【檢查2】是否繳交過 - checkbox格式（第3-5行）")
for idx, student_name in [(3, 'member1'), (4, 'member2'), (5, 'member3')]:
    row = table.rows[idx]
    if len(row.cells) > 9:
        cell = row.cells[9]
        text = cell.text
        has_checkbox_yes = '是' in text
        has_checkbox_no = '否' in text
        has_if_logic = 'if' in text and 'has_submitted' in text

        if has_checkbox_yes and has_checkbox_no and has_if_logic:
            print(f"  ✅ {student_name}: 包含是/否checkbox邏輯")
        else:
            print(f"  ❌ {student_name}: checkbox邏輯不完整")
            print(f"     內容: {text[:100]}")

# 檢查3: 姓名欄位對齊方式
print("\n【檢查3】姓名欄位對齊方式（第3-5行）")
alignments = []
for idx, student_name in [(3, 'member1'), (4, 'member2'), (5, 'member3')]:
    row = table.rows[idx]
    if len(row.cells) > 6:
        cell = row.cells[6]
        if len(cell.paragraphs) > 0:
            alignment = cell.paragraphs[0].alignment
            alignments.append((student_name, alignment))

            # 檢查是否包含姓名變數
            text = cell.paragraphs[0].text
            has_name_var = f"{{{{{student_name}.name}}}}" in text

            status = "✅" if has_name_var else "❌"
            print(f"  {status} {student_name}: 對齊={alignment}, 包含姓名變數={has_name_var}")

# 檢查所有對齊是否一致
if len(set(a[1] for a in alignments)) == 1:
    print(f"  ✅ 所有學生姓名對齊方式一致: {alignments[0][1]}")
else:
    print(f"  ❌ 對齊方式不一致: {alignments}")

# 檢查4: 字型設定（檢查幾個關鍵欄位）
print("\n【檢查4】字型設定 - 新細明體")

key_fields = [
    (1, 1, "計畫名稱"),
    (3, 1, "學生1班級"),
    (3, 6, "學生1姓名"),
    (6, 1, "學習動機"),
]

for row_idx, cell_idx, field_name in key_fields:
    if row_idx < len(table.rows):
        row = table.rows[row_idx]
        if cell_idx < len(row.cells):
            cell = row.cells[cell_idx]
            if len(cell.paragraphs) > 0 and len(cell.paragraphs[0].runs) > 0:
                run = cell.paragraphs[0].runs[0]
                font_name = run.font.name

                # 檢查 eastAsia 字型
                r = run._element
                rPr = r.get_or_add_rPr()
                rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
                east_asia_font = None
                if rFonts is not None:
                    east_asia_font = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')

                if east_asia_font == '新細明體':
                    print(f"  ✅ {field_name}: eastAsia字型=新細明體")
                else:
                    print(f"  ⚠️  {field_name}: eastAsia字型={east_asia_font}")

print("\n" + "=" * 80)
print("驗證完成！")
print("=" * 80)
